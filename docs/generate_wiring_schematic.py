"""
Generate wiring_schematic.docx for the ESP32-S3 Heater Plate Control System.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
import copy
from lxml import etree

OUTPUT = r"C:\Users\alexa\OneDrive\Documents\HeatsinkLabWindTunnel\.claude\worktrees\tender-pike\docs\wiring_schematic.docx"

# ── colour constants ────────────────────────────────────────────────────────
DARK_BLUE   = "1F4E79"   # header bar
LIGHT_BLUE  = "D5E8F0"   # table header shading
WARN_RED    = "FFE6E6"   # warning rows
ALT_GREY    = "F5F5F5"   # alternating row
WHITE       = "FFFFFF"
TEXT_WHITE  = "FFFFFF"
RED_TEXT    = "CC0000"
GREY_TEXT   = "808080"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_shading_and_borders(cell, fill_hex: str, add_borders: bool = False):
    """
    Apply solid background shading and optionally borders to a table cell.
    OOXML CT_TcPr ordering: tcW, gridSpan, hMerge, vMerge, tcBorders, shd, noWrap, ...
    tcBorders MUST appear before shd.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove any existing shd or tcBorders first to avoid duplicates
    for tag in (qn("w:shd"), qn("w:tcBorders")):
        existing = tcPr.find(tag)
        if existing is not None:
            tcPr.remove(existing)

    if add_borders:
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_shading(cell, fill_hex: str):
    """Apply solid background shading to a table cell (no borders)."""
    set_cell_shading_and_borders(cell, fill_hex, add_borders=False)


def set_cell_borders(cell):
    """Not used directly — use set_cell_shading_and_borders instead."""
    pass


def cell_text(cell, text, bold=False, color_hex=None, font_size=9, italic=False):
    """Clear cell and add formatted paragraph."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_section_heading(doc, text, level=1):
    """Add bold section heading in Arial 12pt dark blue."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    return p


def add_note_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("Note: ")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = "Arial"
    return p


def make_table(doc, headers, rows, col_widths_cm=None):
    """
    Build a styled table.
    headers: list of str
    rows: list of list of str; prefix '!' for warning row shading
    col_widths_cm: list of floats in cm
    """
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)

    # Header row
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_shading_and_borders(cell, DARK_BLUE, add_borders=True)
        cell_text(cell, hdr, bold=True, color_hex=TEXT_WHITE, font_size=9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        # detect warning marker
        is_warn = False
        clean_row = []
        for cell_val in row_data:
            if cell_val.startswith("!"):
                is_warn = True
                clean_row.append(cell_val[1:])
            else:
                clean_row.append(cell_val)

        fill = WARN_RED if is_warn else (ALT_GREY if r_idx % 2 == 1 else WHITE)
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(clean_row):
            cell = tr.cells[c_idx]
            set_cell_shading_and_borders(cell, fill, add_borders=True)
            cell_text(cell, val, font_size=9)

    doc.add_paragraph()  # spacing after table
    return table


def set_table_no_split(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        cantSplit.set(qn("w:val"), "true")
        trPr.append(cantSplit)


def add_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # ── Footer ─────────────────────────────────────────────────────────
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        # Build pPr manually to ensure correct element ordering:
        # pStyle -> tabs -> jc (tabs must come before jc in CT_PPrBase)
        p_elem = fp._p
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p_elem.insert(0, pPr)
        # Remove any existing jc or tabs from pPr (keep pStyle if present)
        for tag in (qn("w:jc"), qn("w:tabs")):
            existing = pPr.find(tag)
            if existing is not None:
                pPr.remove(existing)
        # Ensure pStyle is present exactly once (it may already be there from Footer style)
        existing_pStyle = pPr.find(qn("w:pStyle"))
        if existing_pStyle is None:
            pStyle_el = OxmlElement("w:pStyle")
            pStyle_el.set(qn("w:val"), "Footer")
            pPr.insert(0, pStyle_el)
        # tabs (must come before jc in CT_PPrBase schema ordering)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9072")  # right margin at ~16cm
        tabs.append(tab)
        pPr.append(tabs)
        # jc (alignment) — after tabs
        jc_el = OxmlElement("w:jc")
        jc_el.set(qn("w:val"), "left")
        pPr.append(jc_el)

        # Left text
        run_left = fp.add_run("HeatsinkLab Wind Tunnel \u2014 Wiring Schematic v1.0")
        run_left.font.name = "Arial"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor.from_string(GREY_TEXT)

        run_tab = fp.add_run("\t")
        run_tab.font.name = "Arial"
        run_tab.font.size = Pt(8)

        run_page = fp.add_run()
        run_page.font.name = "Arial"
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor.from_string(GREY_TEXT)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = ' PAGE \\* MERGEFORMAT '
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)
        run_page._r.append(fldChar3)

        run_of = fp.add_run(" of ")
        run_of.font.name = "Arial"
        run_of.font.size = Pt(8)
        run_of.font.color.rgb = RGBColor.from_string(GREY_TEXT)

        run_numpages = fp.add_run()
        run_numpages.font.name = "Arial"
        run_numpages.font.size = Pt(8)
        run_numpages.font.color.rgb = RGBColor.from_string(GREY_TEXT)
        fldChar1b = OxmlElement("w:fldChar")
        fldChar1b.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = ' NUMPAGES \\* MERGEFORMAT '
        fldChar2b = OxmlElement("w:fldChar")
        fldChar2b.set(qn("w:fldCharType"), "separate")
        fldChar3b = OxmlElement("w:fldChar")
        fldChar3b.set(qn("w:fldCharType"), "end")
        run_numpages._r.append(fldChar1b)
        run_numpages._r.append(instrText2)
        run_numpages._r.append(fldChar2b)
        run_numpages._r.append(fldChar3b)

        # ── Header (not first page) ─────────────────────────────────────────
        section.different_first_page_header_footer = True
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h = hp.add_run("ESP32-S3 Heater Control \u2014 CONFIDENTIAL INTERNAL DOCUMENT")
        run_h.font.name = "Arial"
        run_h.font.size = Pt(8)
        run_h.italic = True
        run_h.font.color.rgb = RGBColor.from_string(GREY_TEXT)

        # First page header stays blank
        first_hdr = section.first_page_header
        first_hdr.is_linked_to_previous = False


# ── main document builder ────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page setup (A4, 1 inch margins) ──────────────────────────────────────
    from docx.shared import Mm
    for section in doc.sections:
        section.page_width  = Mm(210)
        section.page_height = Mm(297)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════════
    #  TITLE PAGE BLOCK
    # ═══════════════════════════════════════════════════════════════════════

    # Blue title bar paragraph (simulated with shaded table 1x1)
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    title_cell = title_table.cell(0, 0)
    set_cell_shading(title_cell, DARK_BLUE)
    title_cell.width = Inches(6.5)
    tc_para = title_cell.paragraphs[0]
    tc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tc_para.paragraph_format.space_before = Pt(10)
    tc_para.paragraph_format.space_after = Pt(4)
    r_title = tc_para.add_run("ESP32-S3 Heater Plate Control System \u2014 Wiring Schematic")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor.from_string(WHITE)

    subtitle_para = title_cell.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_para.paragraph_format.space_before = Pt(2)
    subtitle_para.paragraph_format.space_after = Pt(10)
    r_sub = subtitle_para.add_run("HeatsinkLab Wind Tunnel \u2014 Breadboard Reproduction Guide")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor.from_string(LIGHT_BLUE)

    doc.add_paragraph()  # spacing

    # Version / date line
    ver_p = doc.add_paragraph()
    r_ver = ver_p.add_run("Version: 1.0  |  Date: 2026-04-04")
    r_ver.font.name = "Arial"
    r_ver.font.size = Pt(10)
    r_ver.font.color.rgb = RGBColor.from_string(GREY_TEXT)
    r_ver.italic = True

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 1 — COMPONENT LIST
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. Component List")

    headers_1 = ["Ref", "Component", "Description", "Interface", "Notes"]
    widths_1   = [1.3, 4.2, 5.5, 3.2, 5.0]
    rows_1 = [
        ["U1", "ESP32-S3-DevKitC-1 (WROOM-1-N16R8)", "Main controller, 16MB Flash, 8MB PSRAM", "USB + GPIO", "Powered by USB; onboard 3.3V LDO"],
        ["U2", "XY-MOS MOSFET Driver Module",          "PWM-controlled MOSFET switch for heater",  "PWM signal", "VIN 5\u201324V (set to 12V); 3.3V logic compatible"],
        ["U3", "INA226 Power Monitor (R100 shunt)",    "Measures heater current, voltage, power",  "I2C @ 0x40", "A0=GND, A1=GND; shunt nominal 100mOhm, calibrated 35.3mOhm"],
        ["U4", "MAX6675 Thermocouple Module",           "K-type thermocouple to SPI digital",       "SPI (bit-bang)", "3.3V supply; CS active-low"],
        ["PS1","Lab Power Supply",                      "12V DC heater supply",                     "Terminal",    "Set to 12.0V DC"],
        ["TC1","K-Type Thermocouple",                  "Temperature probe on heater plate",        "Screw terminals (U4)", "Polarity critical: + to T+, \u2013 to T\u2013"],
        ["HL1","Heater Plate",                          "Resistive heating element",                "Screw terminals (U2)", "OUT+ to HL1 to OUT\u2013"],
    ]
    make_table(doc, headers_1, rows_1, widths_1)

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 2 — POWER RAILS
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. Power Rails")

    headers_2 = ["Rail", "Source", "Voltage", "Consumers", "Notes"]
    widths_2   = [1.8, 3.5, 2.0, 5.5, 6.4]
    rows_2 = [
        ["+12V",   "PS1 (Lab PSU)",     "12.0V DC", "XY-MOS VIN+, INA226 IN+",         "Positive terminal of lab supply; high-current path"],
        ["+3.3V",  "U1 onboard LDO",   "3.3V",     "INA226 VCC, MAX6675 VCC",          "Use DevKit 3V3 pin only; do not add external regulator"],
        ["GND",    "Common node",       "0V",       "All modules + PSU\u2013",           "ALL grounds share one node on breadboard"],
        ["USB 5V", "PC USB port",       "5V",       "U1 only",                           "USB-C on DevKitC-1; powers ESP32 only"],
    ]
    make_table(doc, headers_2, rows_2, widths_2)

    # Warning paragraph
    warn_p = doc.add_paragraph()
    warn_p.paragraph_format.space_before = Pt(4)
    warn_p.paragraph_format.space_after = Pt(8)
    r_warn = warn_p.add_run(
        "IMPORTANT: The 12V power circuit and the 3.3V logic circuit share only a single common GND node. "
        "Never connect the 12V rail to any ESP32-S3 or module logic pin."
    )
    r_warn.bold = True
    r_warn.font.name = "Arial"
    r_warn.font.size = Pt(10)
    r_warn.font.color.rgb = RGBColor.from_string(RED_TEXT)

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 3 — GPIO PIN ASSIGNMENTS
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. GPIO Pin Assignments")

    headers_3 = ["GPIO", "Direction", "Signal Name", "Connected To", "Wire Colour", "Notes"]
    widths_3   = [1.8, 1.8, 2.8, 4.2, 2.2, 6.4]
    rows_3 = [
        ["GPIO 4",   "OUTPUT",  "HEATER_PWM",  "XY-MOS TRIG/PWM",          "Orange", "PWM ch.0; 8-bit; ledc peripheral"],
        ["GPIO 5",   "OUTPUT",  "FAN_PWM",     "(not yet connected)",       "Yellow", "Reserved for future fan"],
        ["GPIO 11",  "INPUT",   "TC_SO (MISO)","MAX6675 SO",                "Blue",   "Soft SPI; read-only"],
        ["GPIO 12",  "OUTPUT",  "TC_CS",       "MAX6675 CS",                "Green",  "Active LOW"],
        ["GPIO 13",  "OUTPUT",  "TC_SCK",      "MAX6675 SCK",               "White",  "Soft SPI clock"],
        ["GPIO 15",  "I/O",     "I2C_SDA",     "INA226 SDA",                "Purple", "4.7k\u03a9 pull-up to 3.3V required"],
        ["GPIO 16",  "OUTPUT",  "I2C_SCL",     "INA226 SCL",                "Grey",   "4.7k\u03a9 pull-up to 3.3V required"],
        ["3V3 pin",  "PWR OUT", "+3.3V",       "INA226 VCC, MAX6675 VCC",   "Red",    "Onboard LDO; ~500mA max"],
        ["GND pin",  "PWR",     "Common GND",  "All module GNDs + PSU\u2013","Black",  "Use multiple GND pins if needed"],
    ]
    make_table(doc, headers_3, rows_3, widths_3)

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 4 — MODULE WIRING DETAIL
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Module-by-Module Wiring")

    # 4.1 XY-MOS
    add_sub_heading(doc, "4.1 XY-MOS MOSFET Driver Module")
    headers_41 = ["XY-MOS Pin", "Connect To", "Wire", "Notes"]
    widths_41  = [2.8, 5.5, 3.2, 7.7]
    rows_41 = [
        ["VIN+",     "INA226 IN\u2013 node (+12V after shunt)", "Red (heavy)",   "High-current; use 22AWG min"],
        ["VIN\u2013","Common GND",                              "Black (heavy)", "High-current return"],
        ["OUT+",     "Heater plate + terminal",                 "Red (heavy)",   "Switched 12V to heater"],
        ["OUT\u2013","Heater plate \u2013 terminal",            "Black (heavy)", "Return; back to GND"],
        ["GND",      "Common GND (logic)",                      "Black",         "Logic ground; must also tie to common"],
        ["TRIG/PWM", "ESP32 GPIO 4",                            "Orange",        "3.3V PWM; no level shifter needed"],
    ]
    make_table(doc, headers_41, rows_41, widths_41)
    add_note_paragraph(doc, "The XY-MOS has two ground connections: power GND (VIN\u2013) and signal GND (GND pin). Both must connect to the common ground node.")

    # 4.2 INA226
    add_sub_heading(doc, "4.2 INA226 Power Monitor Module")
    headers_42 = ["INA226 Pin", "Connect To", "Wire", "Notes"]
    widths_42  = [2.8, 5.5, 3.2, 7.7]
    rows_42 = [
        ["IN+",  "PS1 +12V terminal",      "Red",          "High-side entry; before shunt"],
        ["IN\u2013","XY-MOS VIN+",         "Red",          "After shunt; to load supply"],
        ["VBS",  "Same node as IN\u2013",  "Short jumper", "Bus voltage sense point"],
        ["ALE",  "Not connected",          "\u2014",        "Alert unused in firmware"],
        ["SDA",  "ESP32 GPIO 15",          "Purple",        "4.7k\u03a9 pull-up to 3.3V on breadboard"],
        ["SCL",  "ESP32 GPIO 16",          "Grey",          "4.7k\u03a9 pull-up to 3.3V on breadboard"],
        ["GND",  "Common GND",             "Black",         "Logic ground"],
        ["VCC",  "ESP32 3V3 pin",          "Red (thin)",    "3.3V logic supply"],
    ]
    make_table(doc, headers_42, rows_42, widths_42)
    add_note_paragraph(doc,
        "I2C address 0x40 \u2014 A0 and A1 are GND by default on this module. "
        "Shunt marked R100 (100m\u03a9 nominal) but firmware calibrates to 35.3m\u03a9 \u2014 "
        "do not change without re-measuring.")

    # 4.3 MAX6675
    add_sub_heading(doc, "4.3 MAX6675 Thermocouple Module")
    headers_43 = ["MAX6675 Pin", "Connect To", "Wire", "Notes"]
    widths_43  = [2.8, 5.5, 3.2, 7.7]
    rows_43 = [
        ["GND", "Common GND",           "Black",           "Logic ground"],
        ["VCC", "ESP32 3V3 pin",         "Red",             "3.3V supply"],
        ["SCK", "ESP32 GPIO 13",         "White",           "SPI clock"],
        ["CS",  "ESP32 GPIO 12",         "Green",           "Chip select, active LOW"],
        ["SO",  "ESP32 GPIO 11",         "Blue",            "Data out to ESP32 (MISO)"],
        ["T+",  "Thermocouple + lead",   "Thermocouple wire","Check probe markings for polarity"],
        ["T\u2013","Thermocouple \u2013 lead","Thermocouple wire","Reverse = wrong readings"],
    ]
    make_table(doc, headers_43, rows_43, widths_43)
    add_note_paragraph(doc,
        "MAX6675 uses software SPI (bit-bang). INA226 uses I2C. No bus conflict exists between them.")

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 5 — SYSTEM BLOCK DIAGRAM
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. System Block Diagram")

    ascii_lines = [
        " +---------------------------------------------------------------------------+",
        " |                   12V POWER CIRCUIT (High Current)                       |",
        " |                                                                           |",
        " |  +----------+    +-----------------+    +-------------+    +---------+   |",
        " |  | LAB PSU  |    |  INA226 Module  |    |  XY-MOS     |    | HEATER  |   |",
        " |  |  12V DC  |    |  (R100 shunt)   |    |  MOSFET     |    |  PLATE  |   |",
        " |  |          |    |                 |    |             |    |         |   |",
        " |  |   12V+---+--->| IN+    IN- -----+--->|VIN+   OUT+--+--->|  +      |   |",
        " |  |          |    |  (shunt inside) |    |             |    |         |   |",
        " |  |   GND----+----+-----------------+--->|VIN-   OUT---+--->|  -      |   |",
        " |  +----------+    +-------+---------+    +------+------+    +---------+   |",
        " |                          | VBS,SDA,SCL     GND |                         |",
        " +--------------------------|---------------------|-------------------------+",
        "                            |                     |",
        " +--------------------------|---------------------|-------------------------+",
        " |                   3.3V LOGIC CIRCUIT           |                         |",
        " |                                                |                         |",
        " |  +---------------------------------------------+--------------------+   |",
        " |  |           ESP32-S3-DevKitC-1 (USB powered)                       |   |",
        " |  |                                                                   |   |",
        " |  |   GPIO 4  --------------------------------> XY-MOS TRIG/PWM      |   |",
        " |  |   GPIO 15 (SDA) <------------------------------> INA226 SDA      |   |",
        " |  |   GPIO 16 (SCL) --------------------------------> INA226 SCL      |   |",
        " |  |   GPIO 13 (SCK) --------------------------------> MAX6675 SCK     |   |",
        " |  |   GPIO 12 (CS)  --------------------------------> MAX6675 CS      |   |",
        " |  |   GPIO 11 (SO)  <------------------------------ MAX6675 SO       |   |",
        " |  |   3V3 pin        --------------------------------> INA226 VCC     |   |",
        " |  |   3V3 pin        --------------------------------> MAX6675 VCC    |   |",
        " |  |   GND            --------------------------------> All GNDs       |   |",
        " |  +-------------------------------------------------------------------+   |",
        " +---------------------------------------------------------------------------+",
    ]

    # Use a single-cell table with monospace text to preserve ASCII art
    art_table = doc.add_table(rows=1, cols=1)
    art_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    art_cell = art_table.cell(0, 0)
    set_cell_shading(art_cell, "F8F8F8")
    art_cell.width = Inches(6.5)

    first = True
    for line in ascii_lines:
        if first:
            p = art_cell.paragraphs[0]
            first = False
        else:
            p = art_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(7)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 6 — I2C BUS DETAILS
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "6. I2C Bus Details")

    headers_6 = ["Parameter", "Value"]
    widths_6   = [5.5, 13.7]
    rows_6 = [
        ["Bus Master",          "ESP32-S3 (Arduino Wire library)"],
        ["SDA Pin",             "GPIO 15"],
        ["SCL Pin",             "GPIO 16"],
        ["Pull-up Resistors",   "4.7 k\u03a9 from SDA to 3.3V AND 4.7 k\u03a9 from SCL to 3.3V"],
        ["Bus Speed",           "100 kHz (Wire default)"],
        ["INA226 Address",      "0x40 (A0=GND, A1=GND)"],
        ["Devices on Bus",      "INA226 only (no conflicts)"],
    ]
    make_table(doc, headers_6, rows_6, widths_6)

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 7 — CRITICAL WARNINGS
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "7. Critical Warnings")

    warnings = [
        ("WARNING 1: HIGH VOLTAGE / BURN RISK",
         "The heater plate, OUT+/OUT\u2013 wires, VIN+/VIN\u2013 wires, and INA226 IN+/IN\u2013 carry 12V DC at "
         "several amps. Keep all high-voltage wiring away from ESP32-S3 logic pins at all times."),
        ("WARNING 2: GROUND MUST BE COMMON",
         "The PSU negative terminal MUST share the same GND rail as the ESP32-S3. Separated grounds will "
         "prevent PWM control from working and corrupt INA226 readings."),
        ("WARNING 3: SHUNT CALIBRATION",
         "The INA226 module shunt is marked R100 (nominally 100m\u03a9). Firmware uses setMaxCurrentShunt(2.0f, 0.0353f), "
         "meaning the physical shunt measured ~35.3m\u03a9. Do not alter this value without re-measuring your specific module."),
        ("WARNING 4: THERMOCOUPLE POLARITY",
         "K-type thermocouples are polarity-sensitive. Swapping + and \u2013 will produce inverted or erratic "
         "temperature readings. Verify polarity against the probe\u2019s datasheet or markings before powering on."),
        ("WARNING 5: MOSFET MODULE RATINGS",
         "The XY-MOS module has a maximum rated voltage and current. Do not connect heater loads that exceed these limits."),
    ]

    for label, body in warnings:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        # Shaded background via a tiny table
        warn_t = doc.add_table(rows=1, cols=1)
        warn_t.alignment = WD_TABLE_ALIGNMENT.LEFT
        wc = warn_t.cell(0, 0)
        set_cell_shading_and_borders(wc, "FFE6E6", add_borders=True)
        wc.width = Inches(6.5)
        wp = wc.paragraphs[0]
        wp.paragraph_format.space_before = Pt(4)
        wp.paragraph_format.space_after = Pt(4)
        r_label = wp.add_run(label + "  ")
        r_label.bold = True
        r_label.font.name = "Arial"
        r_label.font.size = Pt(9)
        r_label.font.color.rgb = RGBColor.from_string(RED_TEXT)
        r_body = wp.add_run(body)
        r_body.font.name = "Arial"
        r_body.font.size = Pt(9)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 8 — BREADBOARD CHECKLIST
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "8. Breadboard Reproduction Checklist")

    checklist = [
        "Connect all GND nodes: PSU\u2013, ESP32 GND, INA226 GND, MAX6675 GND, XY-MOS GND \u2192 single common rail",
        "Connect ESP32 3V3 pin to INA226 VCC and MAX6675 VCC",
        "Install 4.7 k\u03a9 pull-up resistors: SDA (GPIO 15) to 3.3V and SCL (GPIO 16) to 3.3V",
        "Wire INA226 IN+ to PSU +12V; INA226 IN\u2013 to XY-MOS VIN+",
        "Connect INA226 VBS to the same node as IN\u2013",
        "Connect XY-MOS VIN\u2013 to GND; also connect XY-MOS signal GND pin to GND",
        "Wire ESP32 GPIO 4 to XY-MOS TRIG/PWM",
        "Wire MAX6675 SPI: SCK\u2192GPIO 13, CS\u2192GPIO 12, SO\u2192GPIO 11",
        "Wire MAX6675 VCC\u21923V3, GND\u2192GND",
        "Attach K-type thermocouple to MAX6675 T+/T\u2013 with correct polarity",
        "Connect heater plate to XY-MOS OUT+ and OUT\u2013",
        "Power ESP32 via USB to PC; power heater circuit via lab PSU at 12V",
        'Open serial monitor and verify "INA226 OK" before enabling heater output',
    ]

    for i, item in enumerate(checklist, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        # Override automatic numbering label by clearing and re-adding
        # (python-docx List Number style handles numbering automatically)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    #  SECTION 9 — FIRMWARE PIN CONSTANTS
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "9. Firmware Pin Constants Reference")

    headers_9 = ["Constant", "Value", "Purpose"]
    widths_9   = [4.5, 3.5, 11.2]
    rows_9 = [
        ["heaterMosfetPin", "GPIO 4",       "PWM output to XY-MOS TRIG"],
        ["fanPwmPin",        "GPIO 5",       "Fan PWM (reserved, not wired)"],
        ["thermocoupleSCK",  "GPIO 13",      "MAX6675 clock"],
        ["thermocoupleCS",   "GPIO 12",      "MAX6675 chip select"],
        ["thermocoupleSO",   "GPIO 11",      "MAX6675 data (MISO)"],
        ["i2cSdaPin",        "GPIO 15",      "INA226 SDA"],
        ["i2cSclPin",        "GPIO 16",      "INA226 SCL"],
        ["INA226 address",   "0x40",         "I2C address (A0=A1=GND)"],
        ["Shunt value",      "0.0353 Ohm",   "Calibrated (see Warning 3)"],
        ["Max current",      "2.0 A",        "INA226 safety limit"],
        ["Averaging",        "1024 samples", "Noise rejection for PWM switching"],
        ["PWM channel",      "0",            "ledc peripheral channel"],
    ]
    make_table(doc, headers_9, rows_9, widths_9)

    # ═══════════════════════════════════════════════════════════════════════
    #  HEADERS & FOOTERS
    # ═══════════════════════════════════════════════════════════════════════
    add_header_footer(doc)

    doc.save(OUTPUT)

    # Post-process: fix w:zoom element in settings.xml
    # python-docx emits <w:zoom w:val="bestFit"/> but the schema requires w:percent attribute
    # Fix by replacing with a valid zoom element
    import zipfile, shutil, os, tempfile
    tmp_path = OUTPUT + ".tmp"
    with zipfile.ZipFile(OUTPUT, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/settings.xml':
                    text = data.decode('utf-8')
                    # Replace zoom bestFit (no percent attr) with 100% zoom
                    text = text.replace(
                        '<w:zoom w:val="bestFit"/>',
                        '<w:zoom w:percent="100"/>'
                    )
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    os.replace(tmp_path, OUTPUT)

    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
